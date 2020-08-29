# V-Indexer - A Project by Ananth Gottimukala aka she11z
# Special thanks to @Anirban Dey and @Scott Hughes
# This is a tool alternative to SANS Voltaire
# Please go through Readme.txt file before using

import os
import csv
import sys
import docx
import xlrd
import getpass
import operator
import optparse
from docx.shared import Pt
from pyfiglet import Figlet
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def banner():

    custom_fig = Figlet(font='standard')
    print(custom_fig.renderText('she11z V-Indexer'))
    print('\n\t\t\t\t\t\t\t\t\t(SANS Index Prettifier)')
    print('\nAuthor: Ananth Gottimukala (she11z)')
    print('GitHub: https://github.com/ananth-she11z')
    print('LinkedIn: https://www.linkedin.com/in/ananth-she11z')
    print('\nSpecial thanks to @Anirban Dey and @Scott Hughes')
    print('-------------------------------------------------------------------------------------------------------')

parser = optparse.OptionParser(usage=banner(), version='\nshe11z V-Indexer 11.1')
parser.add_option('-k', dest='keyword_columns', type=int, help='Enter number of columns you have with keywords in your index')
parser.add_option('-f', dest='index_filename', help='Index filename (.xlsx)')
parser.add_option('-c', dest='course_name', default='SANS', help='Enter which course you are preparing for (Eg: GDAT, GCIH) Default set to "SANS"')
parser.add_option('-s', dest='sheet_name', help='Please specify which sheet to process. (Either -s <SheetName> for specific sheet OR -s <all/ALL> to process all available sheets (Make sure all sheets are in the same format with similar column structure)')
parser.add_option('-a', dest='case', help='Keywords in Upper/Lower/Capitalize case (Eg: -a lower/upper/capitalize)')
(options, arguments) = parser.parse_args()

def usage():

    print('\nPlease use --help or -h option for usage help')
    print('\nExample Usage: python ' + sys.argv[0] + ' -k 2 -f myindex.xlsx -c gdat -s all')
    print('Example Usage: python ' + sys.argv[0] + ' -k 2 -f "My Index.xlsx" -c GDAT -s "Book 1"')
    print('\nNOTE: If you have multiple book/page numbers for one keyword please use "-" as delimiter between page/book numbers (Eg: 115-145)')
    sys.exit()

if not options.index_filename or not options.keyword_columns or not options.sheet_name:
    usage()

myindex_name = options.index_filename
mydoc = docx.Document()
paragraph_format = mydoc.styles['Normal'].paragraph_format  # sets paragraph spacing for "after" field
paragraph_format.space_after = Pt(0)
section = mydoc.sections[0] # sets 2 column layout for document
sectPr = section._sectPr
cols = sectPr.xpath('./w:cols')[0]
cols.set(qn('w:num'),'2')
footer_section = mydoc.sections[0]  # sets footer for document
footer = footer_section.footer
footer_text = footer.paragraphs[0]
footer_text.text = '                                Index Prettify tool by Ananth Gottimukala aka "she11z"'
alphabets = ['a','b','c','d','e','f','g','h','i','j','k','l','m','n','o','p','q','r','s','t','u','v','w','x','y','z']   # Code checks for both upper and lower case
non_alphabets = ['0','1','2','3','4','5','6','7','8','9','`','!','@','#','$','%','^','&','*','(',')','-','_','=','+','\"','[',']','{','}','\\','/','|',':',';',',','.','?','~']
username = getpass.getuser()
course_name = options.course_name

def csv_from_excel(sheet):   # converts your .xlsx file to .csv for backend parsing

    wb = xlrd.open_workbook(myindex_name)
    sh = wb.sheet_by_name(sheet)
    your_csv_file = open('index_parsed.csv', 'a', newline='')
    wr = csv.writer(your_csv_file, quoting=csv.QUOTE_ALL)

    for rownum in range(1,sh.nrows):
        wr.writerow(sh.row_values(rownum))
    your_csv_file.close()

def segregate_by_keywords():    # loop through all the columns with keywords

    try:
        n = options.keyword_columns
        i = 0
        file = open('index_parsed.csv', 'r')
        reader = csv.reader(file)
        output_filename = 'index.csv'
        csv_file_output = open(output_filename, 'a', encoding='utf8', newline='')
        csv_writer = csv.writer(csv_file_output)

        for data in reader:
            for k in range(i,int(n)):
                if data[k]:
                    if options.case.upper() == 'UPPER':
                        csv_writer.writerow([data[k].upper().strip(), data[int(n)], data[int(n)+1], data[int(n)+2]])
                    elif options.case.upper() == 'LOWER':
                        csv_writer.writerow([data[k].lower().strip(), data[int(n)], data[int(n)+1], data[int(n)+2]])
                    else:
                        csv_writer.writerow([data[k].lower().strip(), data[int(n)], data[int(n)+1], data[int(n)+2]])

    except IndexError as e:
        print('\n[-] Please double check - Either number of columns are wrong OR all sheets are not same if "-s all/ALL" selected')
        file.close()
        csv_file_output.close()
        os.remove('index_parsed.csv')
        os.remove('index.csv')
        sys.exit()

def sorting():  # will sort the final .csv file alphabelically

    data = csv.reader(open('index.csv'), delimiter=',')
    sortedlist = sorted(data, key=operator.itemgetter(0))

    with open('index_sorted.csv', 'a', newline='') as f:
      fileWriter = csv.writer(f, delimiter=',')

      for row in sortedlist:
          fileWriter.writerow(row)

def alpha_segregate(alpha): # extracts all index having alphabets

    input_filename = 'index_sorted.csv'
    csv_file_input = open(input_filename, 'r')
    csv_reader = csv.reader(csv_file_input)

    for index in csv_reader:
        if index:
            if index[0].startswith(alpha.upper()) or index[0].startswith(alpha.lower()):
                output_filename = alpha + '.csv'
                csv_file_output = open(output_filename, 'a', encoding='utf8')
                csv_writer = csv.writer(csv_file_output)
                csv_writer.writerow([str(index[0]), 'b' + str(index[2].split('.')[0]) + '/' + 'p' + str(index[3].split('.')[0]), str(index[1])])

def non_alpha_segregate(non_alpha): # extracts all non-alpha index

    input_filename = 'index_sorted.csv'
    csv_file_input = open(input_filename, 'r')
    csv_reader = csv.reader(csv_file_input)

    for index in csv_reader:
        if index:
            if index[0].startswith(non_alpha):
                output_filename = 'others.csv'
                csv_file_output = open(output_filename, 'a', encoding='utf8', newline='')
                csv_writer = csv.writer(csv_file_output)
                csv_writer.writerow([str(index[0]), 'b' + str(index[2].split('.')[0]) + '/' + 'p' + str(index[3].split('.')[0]), str(index[1])])

def alpha_document(alpha):  # main method for documenting alphabetical index

    file_name = alpha + '.csv'
    file_input = open(file_name, 'r')
    reader = csv.reader(file_input)
    page_head_name = alpha.upper() + alpha.lower()
    mydoc.add_heading(page_head_name, 0)

    for i in reader:
        if i:
            if options.case.upper() == 'CAPITALIZE':
                mydoc.add_heading(str(i[0].capitalize() + '  [' + str(i[1]) + ']'),3)
                mydoc.add_paragraph().add_run(str(i[2])).italic = True
            else:
                mydoc.add_heading(str(i[0] + '  [' + str(i[1]) + ']'),3)
                mydoc.add_paragraph().add_run(str(i[2])).italic = True

def delete():

    os.remove('index.csv')
    os.remove('index_parsed.csv')
    print('\n[+] Done!! please find "Index_' + options.sheet_name + '.docx" & "Index_' + options.sheet_name + '.csv" files\'s in current directory')
    print('\n[+] Best Luck for exam !! ;)')
    print('-------------------------------------------------------------------------------------------------------')
    csv_index_name = 'Index_' + options.sheet_name + '.csv'
    os.rename('index_sorted.csv', csv_index_name)

def cover_page():   # .docx cover page

    mydoc.add_heading('\n\n\n\n                                   SANS ' + course_name.upper(), 0)
    mydoc.add_heading(username, 2)

def delete_csv_files():  # For deleting all temporary .csv files

    for alpha in alphabets:
        try:
            r_filename = alpha + '.csv'
            os.remove(r_filename)

        except Exception as e:
            pass

def sheet_selection():  # method to check user input for -s option (specific sheet or all?)

    if options.sheet_name.upper() == 'ALL':
        try:
            xls = xlrd.open_workbook(myindex_name, on_demand=True)
            sheets = xls.sheet_names()

        except FileNotFoundError as e:
            print('\n[-] No such file found. Please double check the name of your index file')
            sys.exit()

        for sheet in sheets:
            try:
                csv_from_excel(sheet)

            except xlrd.biffh.XLRDError as e:
                print('\n[-] No such sheet name found. Please double check the name of sheet you specified')
                sys.exit()

            except FileNotFoundError as e:
                print('\n[-] No such file found. Please double check the name of your index file')
                sys.exit()

    else:
        try:
            csv_from_excel(options.sheet_name)

        except xlrd.biffh.XLRDError as e:
            print('\n[-] No such sheet name found. Please double check the name of sheet you specified')
            sys.exit()

def main(): # main method to start processing

    print('\n[+] Processing ...')
    segregate_by_keywords()
    sorting()

    for alpha in alphabets:
        alpha_segregate(alpha)

    for non_alpha in non_alphabets:
        non_alpha_segregate(non_alpha)

    cover_page()

    for alpha in alphabets:
        try:
            mydoc.add_page_break()
            alpha_document(alpha)

        except Exception as e:
            page_head_name = alpha.upper() + alpha.lower()
            mydoc.add_heading(page_head_name, 0)
            mydoc.add_paragraph().add_run('No index was found for this alphabet').font.color.rgb = RGBColor(0,0,153)

    try:
        others_file_name = 'others.csv' # non_alpha_documentation is here ;)
        file = open(others_file_name, 'r')
        reader = csv.reader(file)

        mydoc.add_page_break()
        mydoc.add_heading('Non-Alpha', 0)

        for j in reader:
                if j:
                    mydoc.add_heading(str(j[0] + '   [' + str(j[1]) + ']'),3)
                    mydoc.add_paragraph().add_run(str(j[2])).italic = True

    except Exception as e:
        pass

    delete_csv_files()
    docx_index_name = 'Index_' + options.sheet_name + '.docx'
    mydoc.save(docx_index_name)

    try:
        file.close()
        os.remove('others.csv')

    except Exception as e:
        pass

if __name__ == "__main__":

    sheet_selection()
    main()
    delete()
